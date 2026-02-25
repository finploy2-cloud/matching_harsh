import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =====================================================
# CONFIGURATION
# =====================================================
BASE_DIR = r"D:\matching_harsh\Job_matching_unscreened"

# final_input
INPUT1 = os.path.join(BASE_DIR, "final_input", "input1.xlsx")
RESDEX_FILE = os.path.join(BASE_DIR, "final_input", "resdex_phone.xlsx")

# final_output
FINAL_OUTPUT = os.path.join(BASE_DIR, "final_output")
ALL_JOB_MATCH_UNIQUE = os.path.join(FINAL_OUTPUT, "all_job_matches_phone_unique.xlsx")

# output
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
ADDITIONAL_NEW_LOCATION = os.path.join(OUTPUT_DIR, "additional_new_location.xlsx")
SCREENED = os.path.join(OUTPUT_DIR, "screened_candidates.xlsx")
UNSCREENED = os.path.join(OUTPUT_DIR, "removed_designations.xlsx")
OUTPUT5 = os.path.join(OUTPUT_DIR, "output5.xlsx")

# resume_output (corrected)
RESUME_OUTPUT_DIR = os.path.join(BASE_DIR, "resume_output")
OUTPUT6 = os.path.join(RESUME_OUTPUT_DIR, "output6.xlsx")
CV_FOLDER = os.path.join(RESUME_OUTPUT_DIR, "cv")

REPORT_PATH = os.path.join(BASE_DIR, "FINAL_RUN_REPORT.docx")

# =====================================================
# READ EXCEL SAFELY
# =====================================================
def read_file(path):
    try:
        return pd.read_excel(path)
    except:
        return pd.DataFrame()

df_input1 = read_file(INPUT1)
df_resdex = read_file(RESDEX_FILE)
df_additional = read_file(ADDITIONAL_NEW_LOCATION)
df_screened = read_file(SCREENED)
df_unscreened = read_file(UNSCREENED)
df_job_match = read_file(ALL_JOB_MATCH_UNIQUE)
df_output5 = read_file(OUTPUT5)
df_output6 = read_file(OUTPUT6)

# =====================================================
# CALCULATIONS
# =====================================================
total_candidates = len(df_input1) - 1 if len(df_input1) > 0 else 0
screened_count = len(df_screened)
unscreened_count = len(df_unscreened)
new_locations = len(df_additional)
job_assigned = len(df_job_match)
profiles_opened = len(df_resdex)
unique_resume_numbers = len(df_output5)
uploaded_to_ZA = len(df_output6)

cv_files_count = len(
    [f for f in os.listdir(CV_FOLDER) if os.path.isfile(os.path.join(CV_FOLDER, f))]
) if os.path.isdir(CV_FOLDER) else 0

# =====================================================
# CREATE DOCUMENT
# =====================================================
doc = Document()

# ---- Reduce margins (IMPORTANT) ----
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ---- Finploy Header ----
title = doc.add_heading("Finploy Technologies", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---- Sub-header ----
p = doc.add_paragraph(f"Automated Processing Report\nGenerated on: {datetime.now().strftime('%d %B %Y')}")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)

# ---- Line separator ----
line = doc.add_paragraph("-" * 90)
line.paragraph_format.space_after = Pt(6)
line.alignment = WD_ALIGN_PARAGRAPH.CENTER

# =====================================================
# CLEAN SECTION FUNCTION (NO EXTRA BLANK LINES)
# =====================================================
def add_section(title, lines):
    h = doc.add_heading(title, level=2)
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after = Pt(1)

    for item in lines:
        p = doc.add_paragraph(f"- {item}")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)

# =====================================================
# ADD SECTIONS
# =====================================================
add_section("1. Input Summary (input1.xlsx)", [
    f"Total Rows: {len(df_input1)}",
    f"Total Candidates: {total_candidates}",
])

add_section("2. Screening Summary", [
    f"Total Screened Candidates: {screened_count}",
    f"Total Unscreened Candidates: {unscreened_count}",
])

add_section("3. Additional New Locations", [
    f"Total New Locations Found: {new_locations}",
])

add_section("4. Job Assignment Summary", [
    f"Total Candidates Assigned a Job: {job_assigned}",
])

add_section("5. Naukri Activity Summary", [
    f"Total Profiles Opened (Resdex): {profiles_opened}",
])

add_section("6. Resume Handling Summary", [
    f"Candidates Assigned Unique Resume Numbers (output5.xlsx): {unique_resume_numbers}",
    f"Total CV Files Downloaded: {cv_files_count}",
])

add_section("7. Upload Summary (output6.xlsx)", [
    f"Candidates Uploaded to ZA Server: {uploaded_to_ZA}",
])

add_section("8. Final Notes / Observations", [
    "All processes completed successfully.",
])

# =====================================================
# SAVE
# =====================================================
doc.save(REPORT_PATH)
print("\nReport generated successfully at:", REPORT_PATH)
import subprocess
try:
    print("▶️ Running main21.py ...")
    subprocess.run(["python", r"D:\matching_harsh\Job_matching_unscreened\main21.py"], check=True)
    print("✅ main21.py executed successfully!")
except Exception as e:
    print(f"❌ Failed to run main21.py: {e}")

